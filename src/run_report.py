#!/usr/bin/python

import docx
import csv
from chatgpt_caller import ChatGPTCaller
from html_downloader import Downloader
from bing_searcher import BingSearch


def process_question(question: str, prompt: str, limit: int=0, retry: int=5):
    """Process the given question by searching from Bing, download from URL, then put
    downloaded contents into ChatGPT to parse into a final readable format.

    Args:
        question (str): The question to search from Bing.
        prompt (str): The prompt to ChatGPT.
        limit (int): Maximum number of characters in the answer
        retry (int): Retry the operation up to this number of times.
    Returns:
        str: The parsed answer, empty string if there's an error.
    """
    # Step 1: search from web
    print(f"Step 1: Searching from Web: {question}")
    bing = BingSearch()
    results = bing.search(question)
    if results is None:
        return ""
    pages = results["webPages"]["value"]
    snippet_list = [page["snippet"] for page in pages]

    page_text_list = []
    # Step 2: download context from top url
    print("Step 2: Downloading top URL")
    for page in pages[:3]:
        print(f"Step 2: Downloading {page['url']}")
        text_list = Downloader().download_text(page["url"])
        key = page["snippet"].strip()[:10]
        key_text = [t for t in text_list if key in t]
        page_text_list += key_text

    # Step 3: combine contents
    print("Step 3: Calling ChatGPT to generate the report")
    snippet_text = "\n".join(snippet_list)
    page_text = "\n".join(page_text_list)
    limit_text = ""
    if limit > 0:
        limit_text = f"under {limit} words"
    message_content = f"{prompt} based on the following data {limit_text}\n{page_text}\n{snippet_text}"
    message = [
        {"role": "system", "content": "You are a financial analyst."},
        {"role": "user", "content": message_content[:4096]},
    ]
    response = ChatGPTCaller().send_message(message)

    # Filter out invalid ChatGPT responses and retry up to retry times
    forbidden_words = [
        "cannot",
        "not possible",
        "not provided",
        "does not provide",
        "not clear",
    ]
    if any([s in response.lower() for s in forbidden_words]):
        print(f"Invalid answer: {response}, retrying for {retry} times...")
        print(question, prompt, limit)
        if retry > 0:
            response = process_question(question, prompt, limit, retry - 1)
        else:
            response = ""

    return response


questionnaire_filename = "input/questionnaire.csv"
subjects_filename = "input/subjects.csv"
out_filename = "output/report_text.txt"
report_filename = "output/report.docx"

with open(questionnaire_filename, newline="") as questionnaire_csv:
    reader = csv.DictReader(questionnaire_csv, delimiter=",")
    questionnaire_list = [row for row in reader]

with open(subjects_filename, newline="") as subjects_csv:
    reader = csv.DictReader(subjects_csv, delimiter=",")
    subject_list = [row for row in reader]

doc = docx.Document()

with open(out_filename, "w") as f:
    f.write("")

with open(out_filename, "a") as f:
    for subject in subject_list:
        question_subject = subject["Question subject"]
        prompt_subject = subject["Prompt subject"]
        doc.add_heading(question_subject, 0)
        f.write(question_subject)
        f.write(":\n")

        for questionnaire in questionnaire_list:
            question = questionnaire["Question"]
            prompt = questionnaire["Prompt"]
            limit = int(questionnaire["Length limit"])
            answer = process_question(
                question=question.format(prompt_subject),
                prompt=prompt.format(question_subject),
                limit=limit,
            )
            doc.add_heading(question.format(question_subject), 3)
            doc.add_paragraph(answer)
            doc.save(report_filename)
            f.write(question.format(question_subject))
            f.write("\n")
            f.write(answer)
            f.write("\n\n")

        f.write("")
        f.write("\n\n\n")

doc.save(report_filename)
